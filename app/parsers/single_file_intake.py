from typing import Dict, Any, List, Optional
import io
import pandas as pd
from docx import Document
from app.utils.diagnostics import DiagnosticContext
from .procurement_pdf import parse_procurement_pdf, _extract_text_safe
import re
import pdfplumber
from app.services.insights import compute_procurement_insights, DEFAULT_BASKET

# Column synonym maps for tolerant CSV/Excel intake
MAP = {
  "budget": {"budget","budget_sar","planned","plan","budget_value","planned_sar"},
  "actual": {"actual","actuals","spent","spend","actual_sar","cost_to_date","ctd"},
  "label": {"label","item","description","cost_code","category","project_id","name"}
}

def _map_cols(df: pd.DataFrame) -> pd.DataFrame:
  lower = {c: c.strip().lower() for c in df.columns}
  rename: Dict[str, str] = {}
  used = {c.lower() for c in df.columns}
  for std, alts in MAP.items():
    for col, low in lower.items():
      if low in alts and std not in used:
        rename[col] = std
        used.add(std)
  if rename:
    df = df.rename(columns=rename)
  return df

def _has_budget_actual(df: pd.DataFrame) -> bool:
  cols = {c.lower() for c in df.columns}
  return bool(MAP["budget"] & cols) and bool(MAP["actual"] & cols)

def _strip_num(x: Any) -> Optional[float]:
  try:
    if x is None:
      return None
    s = str(x)
    s = re.sub(r"[^\d\.\-]", "", s).strip()
    return float(s) if s not in ("", "-", ".", "-.") else None
  except Exception:
    return None

def _emit_variance_rows(df: pd.DataFrame) -> List[Dict[str, Any]]:
  """Given a DF that already has 'budget' and 'actual' standardized, emit variance rows."""
  df = df.copy()
  label_col = next((c for c in ["label","item","description","cost_code","category","project_id"] if c in df.columns), None)
  bcol = next((c for c in ["budget","budget_sar"] if c in df.columns), None)
  acol = next((c for c in ["actual","actuals","actual_sar","spent","spend","cost_to_date","ctd"] if c in df.columns), None)
  if not (bcol and acol):
    return []
  df[bcol] = df[bcol].apply(_strip_num)
  df[acol] = df[acol].apply(_strip_num)
  out: List[Dict[str, Any]] = []
  for _, r in df.iterrows():
    b = r.get(bcol)
    a = r.get(acol)
    if b is None and a is None:
      continue
    try:
      out.append({
        "label": str(r.get(label_col) or r.get("item") or r.get("description") or "Line"),
        "budget_sar": b,
        "actual_sar": a,
        "variance_sar": (a - b) if (a is not None and b is not None) else None,
      })
    except Exception:
      continue
  return out

def parse_single_file(filename: str, data: bytes) -> Dict[str, Any]:
  name = (filename or "").lower()
  with DiagnosticContext(file_name=filename, file_size=len(data)) as diag:
    diag.step("start", filename=name)
    if name.endswith(".pdf"):
      diag.step("parse_pdf_start")
      text = _extract_text_safe(data, diag=diag)
      diag.step("pdf_text_extracted", chars=len(text))
      if not text.strip():
        diag.warn("empty_pdf_text")
      lt = text.lower()
      if "budget" in lt and "actual" in lt:
        variance_rows: List[Dict[str, Any]] = []
        try:
          with pdfplumber.open(io.BytesIO(data)) as pdf:
            tables_scanned = 0
            for page in pdf.pages:
              tbls = page.extract_tables() or []
              for t in tbls:
                tables_scanned += 1
                try:
                  df = pd.DataFrame(t[1:], columns=[str(c).strip() for c in t[0]])
                except Exception:
                  continue
                df = _map_cols(df)
                if _has_budget_actual(df):
                  variance_rows.extend(_emit_variance_rows(df))
            diag.step("pdf_tables_scanned", tables=int(tables_scanned), variance_rows=int(len(variance_rows)))
        except Exception as e:
          diag.warn("pdf_table_scan_failed", error=str(e))

        if variance_rows:
          diag.step("mode_variance_pdf", rows=int(len(variance_rows)))
          return {"variance_items": variance_rows, "diagnostics": diag.to_dict()}

      # Procurement path (default)
      ps = parse_procurement_pdf(data, diag=diag, text=text)
      items = ps.get("items") or []
      diag.step("parse_pdf_success", items=len(items))
      needs_llm = (not items) or any(
        (i.get("qty") is None or i.get("unit_price_sar") is None or i.get("amount_sar") is None)
        for i in items
      )
      if needs_llm:
        from app.llm.extract_from_text import extract_items_via_llm
        llm_items: List[Dict[str, Any]] = []
        if text.strip():
          diag.step("llm_fallback_start")
          try:
            llm_items = extract_items_via_llm(text[:4000])
          except Exception as e:
            diag.warn("llm_failed", error=str(e))
        else:
          diag.warn("llm_skipped_no_text")
        if llm_items:
          diag.step("llm_fallback_success", items=len(llm_items))
          if items:
            for base, extra in zip(items, llm_items):
              if base.get("qty") is None and extra.get("qty") is not None:
                base["qty"] = extra.get("qty")
              if base.get("unit_price_sar") is None and extra.get("unit_price_sar") is not None:
                base["unit_price_sar"] = extra.get("unit_price_sar")
              if base.get("amount_sar") is None and extra.get("amount_sar") is not None:
                base["amount_sar"] = extra.get("amount_sar")
              if base.get("description") is None and extra.get("description"):
                base["description"] = extra.get("description")
          else:
            items = [
              {
                "item_code": it.get("co_id"),
                "description": it.get("description"),
                "qty": it.get("qty"),
                "unit_price_sar": it.get("unit_price_sar"),
                "amount_sar": it.get("amount_sar"),
                "vendor_name": ps.get("meta", {}).get("vendor_name"),
                "doc_date": ps.get("meta", {}).get("doc_date"),
                "source": "uploaded_file_llm",
              }
              for it in llm_items
            ]
            ps["items"] = items

      if not items:
        txt = text[:2000]
        if txt:
          items = [{
            "item_code": None,
            "description": txt,
            "qty": None,
            "unit_price_sar": None,
            "amount_sar": None,
            "vendor_name": None,
            "doc_date": None,
            "source": "uploaded_file",
          }]
          ps = {"items": items, "meta": ps.get("meta", {})}

      analysis = compute_procurement_insights(items, basket=DEFAULT_BASKET)
      return {
        "procurement_summary": ps,
        "analysis": analysis,
        "economic_analysis": analysis,
        "insights": analysis,
        "diagnostics": diag.to_dict(),
      }
    if name.endswith(".docx"):
      diag.step("parse_docx_start")
      doc = Document(io.BytesIO(data))
      text = "\n".join(p.text for p in doc.paragraphs)
      diag.step("parse_docx_success", paragraphs=len(doc.paragraphs))
      ps = {"items":[{"item_code": None, "description": text[:2000], "qty": None, "unit_price_sar": None, "amount_sar": None, "vendor_name": None, "doc_date": None, "source":"uploaded_file"}], "meta":{}}
      analysis = compute_procurement_insights(ps.get("items", []), basket=DEFAULT_BASKET)
      return {
        "procurement_summary": ps,
        "analysis": analysis,
        "economic_analysis": analysis,
        "insights": analysis,
        "diagnostics": diag.to_dict(),
        }
    if name.endswith(".csv"):
      diag.step("parse_csv_start")
      try:
        df = pd.read_csv(io.BytesIO(data))
        diag.step("parse_csv_success", rows=int(df.shape[0]), cols=int(df.shape[1]))
      except Exception as e:
        diag.error("read_csv_failed", e)
        return {"error": "failed_to_read_csv", "diagnostics": diag.to_dict()}
      df = _map_cols(df)
    elif name.endswith(".xlsx") or name.endswith(".xls"):
      diag.step("parse_excel_start")
      try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
        diag.step("parse_excel_success", sheets=list(sheets.keys()))
      except Exception as e:
        diag.error("read_excel_failed", e)
        return {"error": "failed_to_read_excel", "diagnostics": diag.to_dict()}
      frames: List[pd.DataFrame] = []
      for sn, sh in sheets.items():
        mapped = _map_cols(sh)
        has_ba = _has_budget_actual(mapped)
        diag.step("sheet_loaded", sheet=sn, rows=int(sh.shape[0]), cols=int(sh.shape[1]), has_budget_actual=has_ba)
        if has_ba:
          frames.append(mapped)
      if frames:
        df = pd.concat(frames, ignore_index=True)
      else:
        first = next(iter(sheets.values()), pd.DataFrame())
        df = _map_cols(first)
    else:
      diag.step("parse_text_start")
      text = data.decode("utf-8", errors="ignore")
      diag.step("parse_text_success", chars=len(text))
      ps = {"items":[{"item_code": None, "description": text[:2000], "qty": None, "unit_price_sar": None, "amount_sar": None, "vendor_name": None, "doc_date": None, "source":"uploaded_file"}], "meta":{}}
      analysis = compute_procurement_insights(ps.get("items", []), basket=DEFAULT_BASKET)
      return {
        "procurement_summary": ps,
        "analysis": analysis,
        "economic_analysis": analysis,
        "insights": analysis,
        "diagnostics": diag.to_dict(),
        }

    df = _map_cols(df)
    diag.step("columns_mapped", columns=list(df.columns))
    if _has_budget_actual(df):
      rows = _emit_variance_rows(df)
      diag.step("mode_variance", rows=int(len(rows)))
      return {"variance_items": rows, "diagnostics": diag.to_dict()}
    else:
      diag.step("mode_procurement_summary")
      items: List[Dict[str,Any]] = []
      for _, r in df.head(50).iterrows():
        desc = " ".join(str(v) for v in r.to_dict().values() if pd.notna(v))[:2000]
        items.append({"item_code": None, "description": desc, "qty": None, "unit_price_sar": None, "amount_sar": None, "vendor_name": None, "doc_date": None, "source":"uploaded_file"})
      analysis = compute_procurement_insights(items, basket=DEFAULT_BASKET)
      return {
        "procurement_summary": {"items": items, "meta": {}},
        "analysis": analysis,
        "economic_analysis": analysis,
        "insights": analysis,
        "diagnostics": diag.to_dict(),
      }
